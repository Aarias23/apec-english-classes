from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parents[1] / "Units_1_and_2_Comprehensive_Exam.docx"

NAVY = "17365D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2937"
MUTED = "5F6B7A"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
BORDER = "B8C2CC"


def set_run_font(run, name="Calibri", size=11, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    total = sum(widths_dxa)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            width = widths_dxa[i]
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    hdr = OxmlElement("w:tblHeader")
    hdr.set(qn("w:val"), "true")
    tr_pr.append(hdr)


def set_paragraph_border(paragraph, side="bottom", color=BORDER, size=6, space=1):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    border = p_bdr.find(qn(f"w:{side}"))
    if border is None:
        border = OxmlElement(f"w:{side}")
        p_bdr.append(border)
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), str(space))
    border.set(qn("w:color"), color)


def shade_paragraph(paragraph, fill=CALLOUT):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "6")
        border.set(qn("w:color"), "D7DEE7")
        borders.append(border)
    p_pr.append(borders)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    return run


def define_numbering(doc, num_fmt="decimal", lvl_text="%1.", left=540, hanging=270):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids or [-1]) + 1
    num_id = max(num_ids or [0]) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), num_fmt)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), str(left))
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(left))
    ind.set(qn("w:hanging"), str(hanging))
    p_pr.extend([tabs, ind])
    lvl.extend([start, fmt, text, suff, p_pr])
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])


def keep_with_next(paragraph, value=True):
    paragraph.paragraph_format.keep_with_next = value


def add_heading(doc, text, level=1, kicker=None, force_page=False):
    if force_page:
        doc.add_page_break()
    if kicker:
        p = doc.add_paragraph(style="Kicker")
        p.add_run(kicker.upper())
        keep_with_next(p)
    p = doc.add_paragraph(text, style=f"Heading {level}")
    keep_with_next(p)
    return p


def add_answer_lines(doc, count=2, height=18):
    for _ in range(count):
        p = doc.add_paragraph(style="Answer Line")
        p.paragraph_format.space_after = Pt(height)
        set_paragraph_border(p, color="AAB4C0", size=4, space=2)


def add_question(doc, text, num_id, points=None):
    p = doc.add_paragraph(style="Question")
    apply_num(p, num_id)
    r = p.add_run(text)
    set_run_font(r, size=11, bold=True)
    if points:
        r = p.add_run(f"  [{points} points]")
        set_run_font(r, size=9.5, color=MUTED, bold=True)
    keep_with_next(p)
    return p


def add_options(doc, options):
    option_num = define_numbering(doc, "lowerLetter", "%1)", left=720, hanging=300)
    for option in options:
        p = doc.add_paragraph(style="Option")
        apply_num(p, option_num)
        p.add_run(option)


def add_callout(doc, text, label=None):
    p = doc.add_paragraph(style="Passage")
    if label:
        r = p.add_run(label + " ")
        set_run_font(r, size=10.5, color=DARK_BLUE, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)
    shade_paragraph(p)
    return p


def add_scoring_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    headers = ("Criterion", "Full-credit evidence", "Points")
    for cell, text in zip(hdr, headers):
        shade_cell(cell, NAVY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, size=9.5, color=WHITE, bold=True)
    set_repeat_table_header(table.rows[0])
    for criterion, evidence, points in rows:
        cells = table.add_row().cells
        for i, text in enumerate((criterion, evidence, str(points))):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 2 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(text)
            set_run_font(r, size=9.25, color=INK, bold=(i == 0))
    set_table_geometry(table, [1900, 6560, 900])
    return table


def setup_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    custom = {
        "Exam Title": (24, NAVY, True, False, 0, 6, 1.0),
        "Exam Subtitle": (13, MUTED, False, False, 0, 18, 1.0),
        "Kicker": (9, BLUE, True, False, 0, 3, 1.0),
        "Question": (11, INK, False, False, 4, 4, 1.15),
        "Option": (10.5, INK, False, False, 0, 2, 1.10),
        "Passage": (10.5, INK, False, False, 4, 10, 1.20),
        "Answer Line": (11, INK, False, False, 0, 8, 1.0),
        "Fine Print": (8.5, MUTED, False, False, 0, 3, 1.0),
    }
    for name, (size, color, bold, italic, before, after, spacing) in custom.items():
        if name not in styles:
            style = styles.add_style(name, 1)
        else:
            style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = bold
        style.font.italic = italic
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = spacing


def setup_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.70)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("LINK IT! 2  |  UNITS 1–2 EXAM")
    set_run_font(r, size=8.5, color=MUTED, bold=True)
    set_paragraph_border(p, color="D7DEE7", size=4, space=4)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Page ")
    set_run_font(r, size=8.5, color=MUTED)
    set_run_font(add_field(p, "PAGE"), size=8.5, color=MUTED)
    r = p.add_run("  •  Units 1 and 2 Comprehensive Exam")
    set_run_font(r, size=8.5, color=MUTED)


def add_front_matter(doc):
    p = doc.add_paragraph(style="Kicker")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("ACADEMIC ASSESSMENT  •  LINK IT! 2")
    p = doc.add_paragraph(style="Exam Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Units 1 and 2 Comprehensive Exam")
    p = doc.add_paragraph(style="Exam Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Unit 1: Making Plans  |  Unit 2: On Vacation")

    info = doc.add_table(rows=2, cols=3)
    info.style = "Table Grid"
    labels = (("Student name", "Student ID", "Section"), ("Date", "Instructor", "Final score / 100"))
    for row, row_labels in zip(info.rows, labels):
        for cell, label in zip(row.cells, row_labels):
            shade_cell(cell, LIGHT_GRAY)
            p = cell.paragraphs[0]
            r = p.add_run(label.upper())
            set_run_font(r, size=8, color=MUTED, bold=True)
            p2 = cell.add_paragraph()
            set_paragraph_border(p2, color="8E99A6", size=4, space=1)
    set_table_geometry(info, [3920, 2720, 2720])

    doc.add_paragraph()
    metrics = doc.add_table(rows=1, cols=3)
    metrics.style = "Table Grid"
    for cell, (label, value) in zip(metrics.rows[0].cells, (("TIME", "100 minutes"), ("TOTAL", "100 points"), ("SCOPE", "Units 1–2"))):
        shade_cell(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label + "\n")
        set_run_font(r, size=8, color=BLUE, bold=True)
        r = p.add_run(value)
        set_run_font(r, size=11, color=NAVY, bold=True)
    set_table_geometry(metrics, [3120, 3120, 3120])

    add_heading(doc, "Instructions", level=1, kicker="Read before you begin")
    instruction_num = define_numbering(doc, "decimal", "%1.", left=540, hanging=270)
    instructions = [
        "Answer every question. Read each prompt completely before responding.",
        "For Part A, select one best answer. Mark your choice clearly.",
        "For Parts B and C, write in complete, understandable English unless the prompt requests a word or phrase.",
        "Use evidence from the passages when a question asks you to explain or justify an answer.",
        "Manage your time: approximately 30 minutes for Part A, 30 minutes for Part B, and 40 minutes for Part C.",
        "Review verb forms, time expressions, spelling, and punctuation before submitting your exam.",
    ]
    for text in instructions:
        p = doc.add_paragraph(style="Normal")
        apply_num(p, instruction_num)
        p.paragraph_format.space_after = Pt(5)
        p.add_run(text)

    add_callout(doc, "Academic integrity: Complete the exam independently. Do not use notes, translation tools, or outside assistance unless your instructor explicitly permits them.", label="Important:")

    add_heading(doc, "Assessment Map", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, ("Section", "Primary evidence", "Points")):
        shade_cell(cell, NAVY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, size=9.5, color=WHITE, bold=True)
    for row in (
        ("Part A", "Vocabulary, grammar, functional language, reading purpose", "40"),
        ("Part B", "Grammar production, reading details, short applied responses", "32"),
        ("Part C", "Planning task and evidence-based vacation review", "28"),
    ):
        cells = table.add_row().cells
        for i, text in enumerate(row):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (0, 2) else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(text)
            set_run_font(r, size=9.5, bold=(i == 0))
    set_table_geometry(table, [1500, 6960, 900])


def add_part_a(doc):
    add_heading(doc, "Part A: Multiple Choice", level=1, kicker="20 questions  •  2 points each  •  40 points", force_page=True)
    p = doc.add_paragraph("Choose the one answer that best completes each sentence or answers each question.")
    p.paragraph_format.space_after = Pt(10)
    qnum = define_numbering(doc, "decimal", "%1.", left=540, hanging=270)
    questions = [
        ("Take an umbrella. It is ___, and there are dark clouds.", ["sunny", "raining", "freezing", "windy"]),
        ("It is sunny, so the best weather-based plan is to ___.", ["go to the park", "stay home because of snow", "cancel every activity", "wear a winter coat indoors"]),
        ("My sister usually ___ her grandparents on Sundays.", ["visit", "visits", "is visiting", "visiting"]),
        ("Listen! The students ___ an invitation dialogue right now.", ["practice", "practices", "are practicing", "is practicing"]),
        ("We ___ outside school at 4:00 this Saturday. The arrangement is already confirmed.", ["meet usually", "are meeting", "meeting", "meets"]),
        ("Which sentence is a polite invitation?", ["You go swimming.", "Would you like to go swimming?", "Why you swim?", "Go swimming yesterday?"]),
        ("Which is the best polite refusal?", ["No.", "I am not like it.", "Sorry, I can't. I'm visiting my family.", "What a bummer! You cannot ask."]),
        ("Choose the connector that gives a reason: “We are staying inside ___ it is raining.”", ["however", "because", "therefore", "for example"]),
        ("“The park is fun. ___, it is too windy today.” Which connector shows contrast?", ["Because", "For example", "However", "Therefore"]),
        ("Which detail makes a future arrangement complete?", ["Only an adjective", "A time and meeting place", "Only a weather word", "A past-time expression"]),
        ("Which sentence correctly describes physical appearance?", ["He is friendly and funny.", "He has short, dark hair.", "He was kind tomorrow.", "He likes vacation reviews."]),
        ("Choose the natural adjective order.", ["She has curly brown long hair.", "She has long, brown, curly hair.", "She is long hair brown.", "She has hair curly long brown."]),
        ("Which word describes personality?", ["bald", "shoulder-length", "friendly", "tall"]),
        ("Which question asks about physical appearance?", ["What is she like?", "Where was she?", "What does she look like?", "When was she there?"]),
        ("My cousins ___ at the beach yesterday.", ["was", "were", "is", "be"]),
        ("I ___ at the hotel last night; I was at my aunt's house.", ["wasn't", "weren't", "isn't", "don't"]),
        ("Which phrase clearly refers to the past?", ["right now", "every Saturday", "two weeks ago", "this weekend"]),
        ("Choose the correctly formed question.", ["Where you were last weekend?", "Where did you were last weekend?", "Where were you last weekend?", "Where was you last weekend?"]),
        ("At age eight, Sofia ___ short hair and a small camera.", ["have", "has", "had", "having"]),
        ("A text with destination photos, persuasive headings, and booking information is most likely ___.", ["a private diary", "a vacation advertisement", "a school rule", "a weather report"]),
    ]
    for stem, options in questions:
        add_question(doc, stem, qnum)
        add_options(doc, options)


def add_part_b(doc):
    add_heading(doc, "Part B: Short Answer and Application", level=1, kicker="8 questions  •  4 points each  •  32 points", force_page=True)
    p = doc.add_paragraph("Respond briefly but completely. Follow the specific requirements in each prompt.")
    p.paragraph_format.space_after = Pt(10)
    qnum = define_numbering(doc, "decimal", "%1.", left=540, hanging=270)

    add_question(doc, "Complete each sentence with the correct form of the verb in parentheses.", qnum, 4)
    for label, text in (
        ("Habit", "Leo usually __________ (play) basketball after class."),
        ("Now", "The students __________ (write) their answers right now."),
        ("Future arrangement", "We __________ (meet) at the library tomorrow."),
        ("Past", "My family __________ (be) in Puerto Plata last July."),
    ):
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Inches(0.35)
        r = p.add_run(label + ": ")
        set_run_font(r, bold=True, color=DARK_BLUE)
        p.add_run(text)

    add_question(doc, "Write a four-line invitation dialogue. Include: a polite invitation, a polite response, a specific time or place, and a confirmed plan or reasonable alternative.", qnum, 4)
    add_answer_lines(doc, 4, 15)

    add_question(doc, "Read the message. Then identify the activity, day/time, place, and backup plan.", qnum, 4)
    add_callout(doc, "Hi, Carla! Would you like to play volleyball on Saturday? We are meeting at 4:30 outside the sports center. The forecast says it may rain. If it rains, we are watching a movie at my house instead. Please let me know!", label="Message:")
    for label in ("Activity:", "Day/time:", "Place:", "Backup plan:"):
        p = doc.add_paragraph(style="Answer Line")
        r = p.add_run(label + " ")
        set_run_font(r, bold=True, color=DARK_BLUE)
        set_paragraph_border(p, color="AAB4C0", size=4, space=2)

    add_question(doc, "Correct the four errors. Rewrite the paragraph accurately.", qnum, 4)
    add_callout(doc, "My brother usually is playing soccer on Sundays. Today, he play at the park. Tomorrow, we meeting him at 5:00. Therefore the weather is rainy, we have an indoor backup plan.", label="Draft:")
    add_answer_lines(doc, 4, 15)

    add_question(doc, "Describe one real or imaginary person in four sentences: two appearance details using is/has, one personality trait, and one respectful identifying detail.", qnum, 4)
    add_answer_lines(doc, 4, 15)

    add_question(doc, "Change each statement into a question. Then write the requested short answer.", qnum, 4)
    p = doc.add_paragraph("a) You were at the beach. (Yes)  →  __________________________________________")
    p.paragraph_format.left_indent = Inches(0.35)
    p = doc.add_paragraph("b) They were in Samana. (Ask where.)  →  ____________________________________")
    p.paragraph_format.left_indent = Inches(0.35)

    add_question(doc, "Complete the four-sentence memory with was, were, wasn't, weren't, or had. Use each blank meaningfully.", qnum, 4)
    add_callout(doc, "Last summer, I ________ in the mountains with my cousins. They ________ very excited. We ________ a small camera, but our parents ________ with us because they were working.", label="Memory:")

    add_question(doc, "Read the vacation review and answer the four questions with evidence from the text.", qnum, 4)
    add_callout(doc, "★★★★☆ Last July, my cousin and I were in the Azores. The islands were green and peaceful, and the people were friendly. We had a small boat tour and saw whales from a respectful distance. One afternoon was rainy, so our hike was disappointing. I still recommend the trip because the scenery and wildlife were incredible.", label="Review:")
    prompts = [
        "Where and when was the writer on vacation?",
        "What positive detail supports the four-star rating?",
        "What was disappointing, and why?",
        "Does the writer recommend the trip? Give the stated reason.",
    ]
    for prompt in prompts:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Inches(0.35)
        r = p.add_run(prompt + " ")
        set_run_font(r, bold=True)
        set_paragraph_border(p, color="AAB4C0", size=4, space=2)


def add_part_c(doc):
    add_heading(doc, "Part C: Essay and Applied Questions", level=1, kicker="2 tasks  •  28 points", force_page=True)
    p = doc.add_paragraph("Plan before you write. Your response must directly address every listed requirement.")
    p.paragraph_format.space_after = Pt(10)
    qnum = define_numbering(doc, "decimal", "%1.", left=540, hanging=270)

    add_question(doc, "Unit 1 Applied Task: Make a Plan That Works", qnum, 12)
    p = doc.add_paragraph("Write an 8–10 line dialogue or one organized paragraph about a group plan. Include all of the following:")
    req_num = define_numbering(doc, "bullet", "•", left=720, hanging=300)
    for text in (
        "a polite invitation and an appropriate response;",
        "a future arrangement using am/is/are + verb-ing;",
        "the activity, time, and place;",
        "weather information and a reason;",
        "one logical connector (because, however, therefore, or for example);",
        "a backup plan that considers the group.",
    ):
        p = doc.add_paragraph(style="Normal")
        apply_num(p, req_num)
        p.paragraph_format.space_after = Pt(2)
        p.add_run(text)
    add_answer_lines(doc, 10, 13)
    add_scoring_table(doc, [
        ("Invitation and response", "Both are polite and context-appropriate.", 2),
        ("Future arrangement", "Accurate am/is/are + verb-ing form.", 2),
        ("Plan details", "Activity, time, and place are clear.", 2),
        ("Weather and reason", "Weather logically affects or supports the plan.", 2),
        ("Connector and backup", "Ideas connect logically; backup is workable.", 2),
        ("Clarity and conventions", "Meaning is clear; spelling and punctuation support reading.", 2),
    ])

    add_question(doc, "Unit 2 Extended Writing: Evidence-Based Vacation Review", qnum, 16)
    p = doc.add_paragraph("Write an 80–100 word review of a real or imaginary vacation. Give it a star rating and include:")
    req_num = define_numbering(doc, "bullet", "•", left=720, hanging=300)
    for text in (
        "where and when the vacation happened and who was there;",
        "at least one respectful appearance or personality description;",
        "accurate past forms of be (was/were/wasn't/weren't) and at least one sentence with had;",
        "two evaluated details (positive, negative, or mixed) supported by specific evidence;",
        "a final recommendation and reason.",
    ):
        p = doc.add_paragraph(style="Normal")
        apply_num(p, req_num)
        p.paragraph_format.space_after = Pt(2)
        p.add_run(text)

    p = doc.add_paragraph()
    r = p.add_run("Star rating:  ☆  ☆  ☆  ☆  ☆     Destination: ")
    set_run_font(r, size=11, bold=True, color=DARK_BLUE)
    set_paragraph_border(p, color="8E99A6", size=4, space=2)
    add_answer_lines(doc, 13, 13)
    p = doc.add_paragraph("Word count: __________")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(p.runs[0], size=9.5, color=MUTED, bold=True)
    add_scoring_table(doc, [
        ("Context", "Place, time, and people are established.", 2),
        ("Description", "Appearance/personality language is accurate and respectful.", 2),
        ("Past be", "Was/were and negative forms agree with their subjects.", 3),
        ("Past have", "Had is used accurately in past context.", 2),
        ("Evidence and evaluation", "Two specific details support the evaluation.", 3),
        ("Verdict", "Rating/recommendation is clear and justified.", 2),
        ("Organization and conventions", "80–100 words; organized, readable, and edited.", 2),
    ])

    add_heading(doc, "Final Check", level=2)
    check_num = define_numbering(doc, "bullet", "☐", left=600, hanging=300)
    for text in (
        "I answered every question.",
        "I checked subject–verb agreement and time expressions.",
        "I supported my opinions with reasons or evidence.",
        "I reviewed spelling, capitalization, and punctuation.",
    ):
        p = doc.add_paragraph(style="Normal")
        apply_num(p, check_num)
        p.paragraph_format.space_after = Pt(3)
        p.add_run(text)


def audit(doc):
    section = doc.sections[0]
    assert round(section.page_width.inches, 2) == 8.50
    assert round(section.page_height.inches, 2) == 11.00
    assert len(doc.paragraphs) > 100
    assert len(doc.tables) >= 5
    text = "\n".join(p.text for p in doc.paragraphs)
    for required in ("Part A: Multiple Choice", "Part B: Short Answer", "Part C: Essay", "100 points", "Unit 1", "Unit 2"):
        assert required in text, required
    assert "ANSWER KEY" not in text.upper()


def build():
    doc = Document()
    setup_styles(doc)
    setup_page(doc)
    core = doc.core_properties
    core.title = "Units 1 and 2 Comprehensive Exam"
    core.subject = "Link It! 2 - Making Plans and On Vacation"
    core.author = "APEC English Program"
    core.keywords = "Unit 1, Unit 2, English exam, Making Plans, On Vacation"
    core.comments = "Student distribution copy"
    add_front_matter(doc)
    add_part_a(doc)
    add_part_b(doc)
    add_part_c(doc)
    audit(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
