from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

import build_units_1_2_exam as base


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "Units_1_and_2_Comprehensive_Exam.docx"
STUDENT = ROOT / "Units_1_and_2_Comprehensive_Exam_Student.docx"
TEACHER = ROOT / "Units_1_and_2_Exam_Teacher_Guide.docx"


def remove_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)


def remove_table(table):
    element = table._element
    element.getparent().remove(element)


def contains_page_break(paragraph):
    return bool(paragraph._p.xpath('.//w:br[@w:type="page"]'))


def compact_student_exam():
    doc = Document(SOURCE)

    # Separate teacher-only material from the student copy.
    for table in list(doc.tables):
        text = " | ".join(cell.text for row in table.rows for cell in row.cells)
        if "Primary evidence" in text or "Full-credit evidence" in text:
            remove_table(table)

    remove_exact = {
        "Assessment Map",
        "Final Check",
        "I answered every question.",
        "I checked subject–verb agreement and time expressions.",
        "I supported my opinions with reasons or evidence.",
        "I reviewed spelling, capitalization, and punctuation.",
    }
    for paragraph in list(doc.paragraphs):
        if paragraph.text in remove_exact or contains_page_break(paragraph):
            remove_paragraph(paragraph)

    # Remove blank spacer paragraphs that do not carry borders, numbering, or fields.
    for paragraph in list(doc.paragraphs):
        if paragraph.text:
            continue
        p_pr = paragraph._p.pPr
        has_border = p_pr is not None and p_pr.find(qn("w:pBdr")) is not None
        has_num = p_pr is not None and p_pr.find(qn("w:numPr")) is not None
        has_field = bool(paragraph._p.xpath(".//w:fldChar"))
        if not (has_border or has_num or has_field) and paragraph.style.name == "Normal":
            remove_paragraph(paragraph)

    # Keep realistic writing room while eliminating oversized answer blocks.
    def trim_answer_block(start_text, keep):
        paragraphs = list(doc.paragraphs)
        start = next(i for i, p in enumerate(paragraphs) if p.text.startswith(start_text))
        candidates = []
        for paragraph in paragraphs[start + 1 :]:
            if paragraph.style.name == "Question":
                break
            if paragraph.style.name == "Answer Line":
                candidates.append(paragraph)
        for paragraph in candidates[keep:]:
            remove_paragraph(paragraph)

    trim_answer_block("Unit 1 Applied Task", 6)
    trim_answer_block("Unit 2 Extended Writing", 9)

    # Compact, consistent paragraph rhythm.
    doc.styles["Normal"].paragraph_format.space_after = Pt(3)
    doc.styles["Normal"].paragraph_format.line_spacing = 1.12
    doc.styles["Question"].paragraph_format.space_before = Pt(3)
    doc.styles["Question"].paragraph_format.space_after = Pt(2)
    doc.styles["Option"].paragraph_format.space_after = Pt(1)
    doc.styles["Passage"].paragraph_format.space_before = Pt(3)
    doc.styles["Passage"].paragraph_format.space_after = Pt(6)
    for paragraph in doc.paragraphs:
        if paragraph.style.name == "Answer Line":
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.0

    for name, before, after in (
        ("Heading 1", 12, 6),
        ("Heading 2", 9, 4),
        ("Heading 3", 7, 3),
        ("Kicker", 0, 2),
    ):
        doc.styles[name].paragraph_format.space_before = Pt(before)
        doc.styles[name].paragraph_format.space_after = Pt(after)

    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)
        section.header_distance = Inches(0.28)
        section.footer_distance = Inches(0.28)

    doc.core_properties.title = "Units 1 and 2 Comprehensive Exam - Student Copy"
    doc.core_properties.comments = "Compact student distribution copy; scoring materials separated"
    doc.save(STUDENT)


def set_cell_text(cell, text, bold=False, centered=False, size=9.5):
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    base.set_run_font(run, size=size, color=base.INK, bold=bold)


def add_key_table(doc):
    answers = [
        (1, "B — raining", 11, "B — has short, dark hair"),
        (2, "A — go to the park", 12, "B — long, brown, curly hair"),
        (3, "B — visits", 13, "C — friendly"),
        (4, "C — are practicing", 14, "C — What does she look like?"),
        (5, "B — are meeting", 15, "B — were"),
        (6, "B — Would you like…?", 16, "A — wasn't"),
        (7, "C — Sorry, I can't…", 17, "C — two weeks ago"),
        (8, "B — because", 18, "C — Where were you…?"),
        (9, "C — However", 19, "C — had"),
        (10, "B — time and place", 20, "B — vacation advertisement"),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, ("Q", "Answer", "Q", "Answer")):
        base.shade_cell(cell, base.NAVY)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        base.set_run_font(run, size=9.5, color=base.WHITE, bold=True)
    base.set_repeat_table_header(table.rows[0])
    for row in answers:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value), bold=i in (0, 2), centered=i in (0, 2))
    base.set_table_geometry(table, [650, 4030, 650, 4030])


def add_short_answer_key(doc):
    rows = [
        ("1", "plays; are writing; are meeting; was", "1 point per correct form", "4"),
        ("2", "Accept equivalent dialogue containing a polite invitation, appropriate response, specific time/place, and confirmation or alternative.", "1 point per required element", "4"),
        ("3", "Volleyball; Saturday at 4:30; outside the sports center; watch a movie at the writer's house.", "1 point per correct detail", "4"),
        ("4", "My brother usually plays soccer on Sundays. Today, he is playing at the park. Tomorrow, we are meeting him at 5:00. Because the weather is rainy, we have an indoor backup plan.", "1 point per corrected error", "4"),
        ("5", "Accept any respectful four-sentence description with two accurate is/has appearance details, one personality trait, and one identifying detail.", "1 point per required element", "4"),
        ("6", "Were you at the beach? Yes, I was. / Where were they?", "2 points per transformation", "4"),
        ("7", "was; were; had; weren't", "1 point per correct form", "4"),
        ("8", "Azores, last July; positive evidence such as the boat tour/whales or scenery/people; the rainy afternoon made the hike disappointing; yes, because the scenery and wildlife were incredible.", "1 point per evidence-based answer", "4"),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, ("Item", "Expected response", "Scoring", "Pts")):
        base.shade_cell(cell, base.NAVY)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        base.set_run_font(run, size=9, color=base.WHITE, bold=True)
    base.set_repeat_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, bold=i == 0, centered=i in (0, 3), size=8.8)
    base.set_table_geometry(table, [600, 5260, 2600, 900])


def build_teacher_guide():
    doc = Document()
    base.setup_styles(doc)
    base.setup_page(doc)
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)
        section.header_distance = Inches(0.3)
        section.footer_distance = Inches(0.3)

    title = doc.add_paragraph(style="Exam Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Units 1 and 2 Exam: Teacher Guide")
    subtitle = doc.add_paragraph(style="Exam Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Answer key, acceptable responses, and scoring criteria")
    base.add_callout(doc, "Use the key as a scoring reference. Accept grammatically accurate equivalent answers that meet the communicative objective. Total: 100 points.", label="Teacher note:")

    base.add_heading(doc, "Part A: Multiple-Choice Key", level=1, kicker="40 points")
    add_key_table(doc)

    base.add_heading(doc, "Part B: Short-Answer Key", level=1, kicker="32 points")
    add_short_answer_key(doc)

    base.add_heading(doc, "Part C: Applied-Writing Criteria", level=1, kicker="28 points")
    base.add_heading(doc, "Task 1 — Make a Plan That Works", level=2)
    base.add_scoring_table(doc, [
        ("Invitation and response", "Both are polite and appropriate to the context.", 2),
        ("Future arrangement", "Accurate am/is/are + verb-ing form.", 2),
        ("Plan details", "Activity, time, and place are clear.", 2),
        ("Weather and reason", "Weather logically affects or supports the plan.", 2),
        ("Connector and backup", "Ideas connect logically; backup is workable and inclusive.", 2),
        ("Clarity and conventions", "Meaning is clear; spelling and punctuation support reading.", 2),
    ])

    base.add_heading(doc, "Task 2 — Evidence-Based Vacation Review", level=2)
    base.add_scoring_table(doc, [
        ("Context", "Place, time, and people are established.", 2),
        ("Description", "Appearance/personality language is accurate and respectful.", 2),
        ("Past be", "Was/were and negative forms agree with their subjects.", 3),
        ("Past have", "Had is used accurately in past context.", 2),
        ("Evidence and evaluation", "Two specific details support the evaluation.", 3),
        ("Verdict", "Rating or recommendation is clear and justified.", 2),
        ("Organization and conventions", "80–100 words; organized, readable, and edited.", 2),
    ])

    base.add_heading(doc, "Scoring Summary", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, ("Section", "Maximum", "Student score")):
        base.shade_cell(cell, base.NAVY)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        base.set_run_font(run, size=9.5, color=base.WHITE, bold=True)
    for row in (("Part A", "40", ""), ("Part B", "32", ""), ("Part C", "28", ""), ("Total", "100", "")):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, bold=row[0] == "Total", centered=True)
    base.set_table_geometry(table, [4560, 2400, 2400])

    doc.styles["Normal"].paragraph_format.space_after = Pt(4)
    doc.core_properties.title = "Units 1 and 2 Exam Teacher Guide"
    doc.core_properties.comments = "Teacher-only answer key and scoring criteria"
    doc.save(TEACHER)


def audit():
    student = Document(STUDENT)
    teacher = Document(TEACHER)
    student_text = "\n".join(p.text for p in student.paragraphs)
    teacher_text = "\n".join(p.text for p in teacher.paragraphs)
    assert "Assessment Map" not in student_text
    assert "Full-credit evidence" not in " ".join(c.text for t in student.tables for r in t.rows for c in r.cells)
    assert "Final Check" not in student_text
    assert len(student.tables) == 2
    assert "Multiple-Choice Key" in teacher_text
    assert "Applied-Writing Criteria" in teacher_text
    assert len(teacher.tables) == 5
    assert sum(1 for p in student.paragraphs if p.style.name == "Answer Line") == 33
    print(STUDENT)
    print(TEACHER)


if __name__ == "__main__":
    compact_student_exam()
    build_teacher_guide()
    audit()
