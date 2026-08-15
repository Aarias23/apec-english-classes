from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

import build_units_1_2_exam as base


ROOT = Path(__file__).resolve().parents[1]
EXAM = ROOT / "Units_1_and_2_Exam_Two_Pages.docx"
GUIDE = ROOT / "Units_1_and_2_Exam_Two_Pages_Teacher_Guide.docx"
WIDE_DXA = 10368  # 7.2 inches; named compact-exam geometry override.


def compact_styles(doc):
    base.setup_styles(doc)
    normal = doc.styles["Normal"]
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.0
    for name, size, before, after in (
        ("Heading 1", 12, 5, 3),
        ("Heading 2", 10.5, 4, 2),
        ("Kicker", 8, 0, 1),
        ("Question", 9.2, 1.5, 0.5),
        ("Passage", 9, 2, 3),
        ("Answer Line", 9, 0, 3),
    ):
        style = doc.styles[name]
        style.font.size = Pt(size)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
    doc.styles["Exam Title"].font.size = Pt(17)
    doc.styles["Exam Title"].paragraph_format.space_after = Pt(2)
    doc.styles["Exam Subtitle"].font.size = Pt(9.5)
    doc.styles["Exam Subtitle"].paragraph_format.space_after = Pt(4)
    if "Inline Options" not in doc.styles:
        style = doc.styles.add_style("Inline Options", WD_STYLE_TYPE.PARAGRAPH)
    else:
        style = doc.styles["Inline Options"]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(8.8)
    style.paragraph_format.left_indent = Inches(0.25)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(1.5)
    style.paragraph_format.line_spacing = 1.0


def compact_page(doc):
    base.setup_page(doc)
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    section.header_distance = Inches(0.22)
    section.footer_distance = Inches(0.22)


def add_info_row(doc):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    labels = ("Name", "Student ID", "Section", "Score / 50")
    for cell, label in zip(table.rows[0].cells, labels):
        base.shade_cell(cell, base.LIGHT_GRAY)
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(label + ": __________________")
        base.set_run_font(run, size=8.5, color=base.INK, bold=True)
        base.set_cell_margins(cell, top=40, bottom=40, start=70, end=70)
    base.set_table_geometry(table, [3200, 2600, 2300, 2268], indent_dxa=70)


def add_section_title(doc, title, details):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(title)
    base.set_run_font(run, size=11.5, color=base.BLUE, bold=True)
    run = paragraph.add_run("  " + details)
    base.set_run_font(run, size=8.5, color=base.MUTED, bold=True)
    base.set_paragraph_border(paragraph, color=base.BLUE, size=6, space=2)


def add_mcq(doc, qnum, stem, options):
    paragraph = doc.add_paragraph(style="Question")
    base.apply_num(paragraph, qnum)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(stem)
    base.set_run_font(run, size=9.2, color=base.INK, bold=True)
    option_paragraph = doc.add_paragraph(style="Inline Options")
    labels = ("A", "B", "C", "D")
    for index, (label, option) in enumerate(zip(labels, options)):
        if index:
            option_paragraph.add_run("     ")
        run = option_paragraph.add_run(label + ". ")
        base.set_run_font(run, size=8.8, color=base.DARK_BLUE, bold=True)
        run = option_paragraph.add_run(option)
        base.set_run_font(run, size=8.8, color=base.INK)


def add_response_line(doc, label="", after=4):
    paragraph = doc.add_paragraph(style="Answer Line")
    if label:
        run = paragraph.add_run(label + " ")
        base.set_run_font(run, size=8.8, color=base.DARK_BLUE, bold=True)
    paragraph.paragraph_format.space_after = Pt(after)
    base.set_paragraph_border(paragraph, color="AAB4C0", size=4, space=1)


def build_exam():
    doc = Document()
    compact_styles(doc)
    compact_page(doc)

    title = doc.add_paragraph(style="Exam Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Units 1 and 2 Exam")
    subtitle = doc.add_paragraph(style="Exam Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Making Plans • On Vacation • 50 points • 60 minutes")
    add_info_row(doc)
    instructions = doc.add_paragraph()
    instructions.paragraph_format.space_before = Pt(2)
    instructions.paragraph_format.space_after = Pt(2)
    run = instructions.add_run("Instructions: ")
    base.set_run_font(run, size=8.8, color=base.DARK_BLUE, bold=True)
    run = instructions.add_run("Answer every question. Choose one answer in Part A and write clearly in Parts B and C.")
    base.set_run_font(run, size=8.8, color=base.INK)

    add_section_title(doc, "Part A: Multiple Choice", "12 questions × 2 points = 24 points")
    qnum = base.define_numbering(doc, "decimal", "%1.", left=360, hanging=240)
    questions = [
        ("Take an umbrella. It is ___ and there are dark clouds.", ("sunny", "raining", "freezing", "dry")),
        ("My sister usually ___ her grandparents on Sundays.", ("visit", "visits", "is visiting", "visited")),
        ("Listen! The students ___ a dialogue right now.", ("practice", "practices", "are practicing", "were practicing")),
        ("We ___ outside school at 4:00 this Saturday. The plan is confirmed.", ("meet usually", "are meeting", "met", "meets")),
        ("Which sentence is a polite invitation?", ("You go swimming.", "Would you like to go swimming?", "Why you swim?", "You swam?")),
        ("We are staying inside ___ it is raining.", ("however", "because", "therefore", "for example")),
        ("Which sentence describes physical appearance?", ("He is very friendly.", "He has short, curly hair.", "He was kind.", "He likes trips.")),
        ("Which question asks about personality?", ("What does she look like?", "What is she like?", "Where was she?", "How old was she?")),
        ("My cousins ___ at the beach yesterday.", ("was", "were", "is", "be")),
        ("Choose the correct past question.", ("Where you were?", "Where did you were?", "Where were you?", "Where was you?")),
        ("At age eight, Sofia ___ short hair and a small camera.", ("have", "has", "had", "having")),
        ("Destination photos, persuasive headings, and booking details usually identify ___.", ("a diary", "a vacation advertisement", "a school rule", "a weather report")),
    ]
    for stem, options in questions:
        add_mcq(doc, qnum, stem, options)

    doc.add_page_break()

    add_section_title(doc, "Part B: Short Answer", "4 questions × 4 points = 16 points")
    qnum = base.define_numbering(doc, "decimal", "%1.", left=360, hanging=240)
    paragraph = doc.add_paragraph(style="Question")
    base.apply_num(paragraph, qnum)
    paragraph.add_run("Complete the verbs: Leo usually ______ (play). We ______ (study) now. We ______ (meet) tomorrow. My family ______ (be) in Puerto Plata last July.")
    paragraph = doc.add_paragraph(style="Question")
    base.apply_num(paragraph, qnum)
    paragraph.add_run("Write a three-line dialogue with a polite invitation, a response, and a confirmed time/place or alternative.")
    add_response_line(doc, "A:")
    add_response_line(doc, "B:")
    add_response_line(doc, "A/B:")
    paragraph = doc.add_paragraph(style="Question")
    base.apply_num(paragraph, qnum)
    paragraph.add_run("Describe one person in four sentences: two appearance details using is/has, one personality trait, and one past sentence using was/were.")
    add_response_line(doc)
    add_response_line(doc)
    add_response_line(doc)
    paragraph = doc.add_paragraph(style="Question")
    base.apply_num(paragraph, qnum)
    paragraph.add_run("Read and answer the four prompts.")
    base.add_callout(doc, "Last July, Elena was in Puerto Plata with her aunt. Her aunt is tall, has short gray hair, and is very friendly. They had a boat trip. Elena recommends the vacation because the beach was beautiful and the people were kind.", label="Review:")
    add_response_line(doc, "Where/when:", 3)
    add_response_line(doc, "Person:", 3)
    add_response_line(doc, "Activity:", 3)
    add_response_line(doc, "Recommendation/reason:", 3)

    add_section_title(doc, "Part C: Applied Writing", "2 tasks × 5 points = 10 points")
    qnum = base.define_numbering(doc, "decimal", "%1.", left=360, hanging=240)
    paragraph = doc.add_paragraph(style="Question")
    base.apply_num(paragraph, qnum)
    paragraph.add_run("Unit 1 — Write 4–5 sentences about a workable plan. Include an invitation, future arrangement, time/place, weather or reason, and a backup plan.")
    for _ in range(4):
        add_response_line(doc, after=3)
    paragraph = doc.add_paragraph(style="Question")
    base.apply_num(paragraph, qnum)
    paragraph.add_run("Unit 2 — Write a 50–60 word vacation review. Include where/when, one person description, was/were, had, one evaluated detail, and a recommendation.")
    for _ in range(5):
        add_response_line(doc, after=3)
    paragraph = doc.add_paragraph("Word count: ________")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_after = Pt(0)
    base.set_run_font(paragraph.runs[0], size=8.5, color=base.MUTED, bold=True)

    doc.core_properties.title = "Units 1 and 2 Exam - Two Page Student Copy"
    doc.core_properties.comments = "Exactly one manual page break; compact two-page design"
    doc.save(EXAM)


def teacher_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, headers):
        base.shade_cell(cell, base.NAVY)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        base.set_run_font(run, size=9, color=base.WHITE, bold=True)
    base.set_repeat_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            paragraph = cells[index].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if index in (0, len(row) - 1) else WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(str(value))
            base.set_run_font(run, size=8.7, color=base.INK, bold=index == 0)
    base.set_table_geometry(table, widths)


def build_guide():
    doc = Document()
    base.setup_styles(doc)
    base.setup_page(doc)
    title = doc.add_paragraph(style="Exam Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Two-Page Exam: Teacher Guide")
    subtitle = doc.add_paragraph(style="Exam Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Answer key and scoring criteria • Units 1 and 2 • 50 points")
    base.add_callout(doc, "Accept equivalent responses that are accurate and meet the stated communicative requirement.", label="Scoring note:")

    base.add_heading(doc, "Part A: Key", level=1, kicker="24 points")
    teacher_table(doc, ("Q", "Answer", "Q", "Answer"), (
        (1, "B — raining", 7, "B — has short, curly hair"),
        (2, "B — visits", 8, "B — What is she like?"),
        (3, "C — are practicing", 9, "B — were"),
        (4, "B — are meeting", 10, "C — Where were you?"),
        (5, "B — Would you like…?", 11, "C — had"),
        (6, "B — because", 12, "B — vacation advertisement"),
    ), (650, 4030, 650, 4030))

    base.add_heading(doc, "Part B: Suggested Answers", level=1, kicker="16 points")
    teacher_table(doc, ("Item", "Expected evidence", "Scoring", "Pts"), (
        (1, "plays; are studying; are meeting; was", "1 point each", 4),
        (2, "Polite invitation; suitable response; time/place; confirmation or alternative", "1 point each", 4),
        (3, "Two accurate is/has appearance details; personality trait; accurate past was/were sentence", "1 point each", 4),
        (4, "Puerto Plata/last July; tall/short gray hair/friendly; boat trip; yes—beautiful beach and kind people", "1 point each", 4),
    ), (650, 4810, 3000, 900))

    base.add_heading(doc, "Part C: Criteria", level=1, kicker="10 points")
    teacher_table(doc, ("Task", "Full-credit requirements", "Pts"), (
        (1, "4–5 clear sentences containing an invitation, accurate future arrangement, time/place, reason or weather, and workable backup plan.", 5),
        (2, "50–60 organized words containing context, person description, accurate was/were and had, evaluated detail, and justified recommendation.", 5),
    ), (850, 7610, 900))
    base.add_heading(doc, "Score Summary", level=2)
    teacher_table(doc, ("Part A", "Part B", "Part C", "Total"), (("/24", "/16", "/10", "/50"),), (2340, 2340, 2340, 2340))
    doc.core_properties.title = "Two-Page Units 1 and 2 Exam Teacher Guide"
    doc.save(GUIDE)


def audit():
    exam = Document(EXAM)
    guide = Document(GUIDE)
    exam_text = "\n".join(p.text for p in exam.paragraphs)
    guide_text = "\n".join(p.text for p in guide.paragraphs)
    page_breaks = sum(1 for p in exam.paragraphs if 'w:br' in p._p.xml and 'page' in p._p.xml)
    assert page_breaks == 1
    assert all(text in exam_text for text in ("Part A: Multiple Choice", "Part B: Short Answer", "Part C: Applied Writing", "50 points"))
    assert "Teacher Guide" in guide_text
    assert len(exam.tables) == 1
    assert len(guide.tables) == 4
    print(EXAM)
    print(GUIDE)


if __name__ == "__main__":
    build_exam()
    build_guide()
    audit()
