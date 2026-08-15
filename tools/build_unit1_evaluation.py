import os
import sys

sys.path.insert(0, r"C:\Users\alexi\AppData\Local\Temp\apec_docx_deps")

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_PATH = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "Unit_1_Making_Plans_Evaluation_Quiz.docx"))

NAVY = "17365D"
BLUE = "2E74B5"
TEAL = "1F6F78"
LIGHT_BLUE = "E8EEF5"
LIGHT_TEAL = "E8F3F4"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
GRID = "B8C2CC"
INK = "1F2937"
WHITE = "FFFFFF"
GOLD = "D99A2B"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color=GRID, size=8, inside=True):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    names = ["top", "left", "bottom", "right"]
    if inside:
        names += ["insideH", "insideV"]
    for name in names:
        edge = borders.find(qn(f"w:{name}"))
        if edge is None:
            edge = OxmlElement(f"w:{name}")
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), str(size))
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), color)


def set_table_geometry(table, widths, indent=120):
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_font(run, name="Calibri", size=11, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, fld_sep, text, fld_end])
    set_font(run, size=9, color=MID_GRAY)


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_bottom_border(paragraph, color="C8CDD3", size=6, space=2):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_section_heading(doc, label, title, points):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(f"{label}  {title}")
    set_font(r, size=16, color=BLUE, bold=True)
    if points != "":
        pr = p.add_run(f"   {points} points")
        set_font(pr, size=9.5, color=TEAL, bold=True)
    return p


def add_instruction(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_font(r, size=10.5, color=MID_GRAY, italic=True)
    return p


def add_question(doc, number, text, options=None, answer_lines=0):
    p = doc.add_paragraph(style="Assessment Question")
    if options or answer_lines:
        p.paragraph_format.keep_with_next = True
    r = p.add_run(f"{number}. ")
    set_font(r, size=10.5, color=NAVY, bold=True)
    r = p.add_run(text)
    set_font(r, size=10.5, color=INK)
    if options:
        p2 = doc.add_paragraph(style="Assessment Options")
        r2 = p2.add_run("    ".join(options))
        set_font(r2, size=10.2, color=INK)
    for _ in range(answer_lines):
        line = doc.add_paragraph()
        line.paragraph_format.space_before = Pt(0)
        line.paragraph_format.space_after = Pt(5)
        line.paragraph_format.line_spacing = 1.0
        line.add_run(" ")
        add_bottom_border(line)


def add_callout(doc, title, body, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table, color=fill, size=6, inside=False)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_font(r, size=10.5, color=NAVY, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_font(r2, size=10.3, color=INK)
    return table


def configure_document(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, before, after, color in (
        ("Heading 1", 16, 18, 10, BLUE),
        ("Heading 2", 13, 14, 7, BLUE),
        ("Heading 3", 12, 10, 5, NAVY),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    styles = doc.styles
    qstyle = styles.add_style("Assessment Question", 1)
    qstyle.base_style = styles["Normal"]
    qstyle.font.name = "Calibri"
    qstyle.font.size = Pt(10.5)
    qstyle.paragraph_format.space_after = Pt(3)
    qstyle.paragraph_format.line_spacing = 1.08
    qstyle.paragraph_format.keep_together = True

    ostyle = styles.add_style("Assessment Options", 1)
    ostyle.base_style = styles["Normal"]
    ostyle.font.name = "Calibri"
    ostyle.font.size = Pt(10.2)
    ostyle.paragraph_format.left_indent = Inches(0.24)
    ostyle.paragraph_format.space_after = Pt(6)
    ostyle.paragraph_format.line_spacing = 1.05
    ostyle.paragraph_format.keep_together = True

    for section in doc.sections:
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hp.paragraph_format.space_after = Pt(0)
        hr = hp.add_run("UNIT 1 EVALUATION  |  MAKING PLANS")
        set_font(hr, size=8.5, color=MID_GRAY, bold=True)

        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        fr = fp.add_run("APEC English  |  Page ")
        set_font(fr, size=9, color=MID_GRAY)
        add_field(fp, "PAGE")
        fr2 = fp.add_run(" of ")
        set_font(fr2, size=9, color=MID_GRAY)
        add_field(fp, "NUMPAGES")


def add_masthead(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run("UNIT 1 • EVALUATION QUIZ")
    set_font(r, size=10, color=TEAL, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Making Plans")
    set_font(r, size=28, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("Weather, verb tenses, invitations, travel plans, reading and writing")
    set_font(r, size=12.5, color=MID_GRAY)

    student = doc.add_table(rows=1, cols=3)
    set_table_geometry(student, [4320, 2520, 2520])
    set_table_borders(student, color=GRID, size=6, inside=True)
    for cell, label in zip(student.rows[0].cells, ("Name: __________________________", "Class: ____________", "Date: ____________")):
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        set_font(r, size=10, color=INK, bold=True)

    doc.add_paragraph().paragraph_format.space_after = Pt(1)

    metrics = doc.add_table(rows=1, cols=3)
    set_table_geometry(metrics, [3120, 3120, 3120])
    set_table_borders(metrics, color=LIGHT_TEAL, size=6, inside=True)
    data = (("TOTAL", "50 points"), ("TIME", "55 minutes"), ("LEVEL", "Unit 1 / A2"))
    for cell, (label, value) in zip(metrics.rows[0].cells, data):
        set_cell_shading(cell, LIGHT_TEAL)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(label)
        set_font(r, size=8.5, color=TEAL, bold=True)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(value)
        set_font(r2, size=11.5, color=NAVY, bold=True)

    add_callout(
        doc,
        "Directions",
        "Read every item carefully. Write clearly. For multiple-choice questions, circle one answer. Complete the writing task in 70–90 words. Dictionaries and translators are not permitted unless your teacher says otherwise.",
        LIGHT_GRAY,
    )


def build_student_pages(doc):
    add_masthead(doc)

    add_section_heading(doc, "A", "WEATHER & MEANING", 8)
    add_instruction(doc, "Use the word bank. Write one weather word for each definition. Each word is used once.")
    add_callout(doc, "Word bank", "cloudy  •  dry  •  foggy  •  stormy  •  warm  •  windy", LIGHT_TEAL)
    for num, definition in enumerate(
        [
            "with many clouds: ____________________",
            "with strong wind: ____________________",
            "pleasant—not hot or cold: ____________________",
            "hard to see because of low cloud: ____________________",
            "with a storm: ____________________",
            "not wet: ____________________",
        ],
        1,
    ):
        add_question(doc, num, definition)
    add_question(doc, 7, "It is rainy and wet. What is the best plan?", ["A. Go swimming outside", "B. Watch a movie inside", "C. Fly a kite"])
    add_question(doc, 8, "The temperature is 4°C. There is strong wind and there are dark clouds. Choose the best report.", ["A. It is warm and sunny.", "B. It is cold, windy and cloudy.", "C. It is hot and dry."])

    add_section_heading(doc, "B", "GRAMMAR IN CONTEXT", 12)
    add_instruction(doc, "Choose the best form. The time expression is your clue.")
    grammar = [
        ("My sister usually ____ music after class.", ["A. listens to", "B. is listening to", "C. listen to"]),
        ("Be quiet! We ____ a video right now.", ["A. record", "B. are recording", "C. records"]),
        ("I ____ my grandparents every Sunday.", ["A. visit", "B. am visiting", "C. visits"]),
        ("What ____ on Friday?", ["A. you do", "B. are you doing", "C. do you doing"]),
        ("They ____ outside school tomorrow at 5:30.", ["A. meet", "B. are meeting", "C. meets"]),
        ("She ____ to the festival this weekend.", ["A. goes", "B. is going", "C. go"]),
        ("Today, I ____ a jacket because it is warm.", ["A. don't wear", "B. am not wearing", "C. doesn't wear"]),
        ("Harry usually ____ tennis at 9:00 on Saturdays.", ["A. plays", "B. is playing", "C. play"]),
    ]
    for num, (stem, choices) in enumerate(grammar, 9):
        add_question(doc, num, stem, choices)
    add_question(doc, 17, "Correct the false plan. Schedule: Saturday, 2:00—study for the science test. Sentence: Harry is playing tennis at 2:00.", answer_lines=2)
    add_question(doc, 18, "Write one question to ask about a future arrangement using when, where, or what time.", answer_lines=1)

    doc.add_page_break()

    add_section_heading(doc, "C", "INVITE & MAKE AN ARRANGEMENT", 8)
    add_instruction(doc, "Complete the dialogue with the expressions in the box. Then identify the final plan.")
    add_callout(doc, "Expression bank", "That sounds great.  •  What a bummer!  •  Would you like to  •  Sorry, I can't.  •  How about  •  Let's meet", LIGHT_TEAL)
    dialogue = [
        "A: Are you doing anything on Saturday?",
        "B: (19) ______________________________ I'm visiting my grandparents.",
        "A: (20) ______________________________ (21) ______________________________ watch a movie on Sunday?",
        "B: (22) ______________________________",
        "A: (23) ______________________________ 4:00?",
        "B: Great. (24) ______________________________ outside the movie theater.",
    ]
    for line in dialogue:
        p = doc.add_paragraph(style="Assessment Question")
        p.paragraph_format.left_indent = Inches(0.18)
        r = p.add_run(line)
        set_font(r, size=10.5, color=INK)
    add_question(doc, 25, "What are they doing, and when? __________________________________________")
    add_question(doc, 26, "Where are they meeting? _________________________________________________")

    add_section_heading(doc, "D", "READING: A WEEKEND BLOG", 10)
    add_instruction(doc, "Read the original blog. Answer in complete sentences when requested.")
    passage = (
        "Hello from Samaná! I am here with my cousin Elena for a family weekend. The town is lively, and the ocean looks amazing. "
        "Right now, Elena is taking photos near the water, and I am writing this post at a small café. It is warm but cloudy today. "
        "Tomorrow morning, we are visiting Los Haitises National Park with my aunt. In the afternoon, we are meeting our cousins at 4:00 outside the town museum. "
        "On Sunday, we are having lunch together before we return home. However, the forecast says it may be stormy tomorrow. If the boat tour is canceled, we are visiting the museum instead. "
        "I think our weekend will be wonderful because we have a clear backup plan."
    )
    add_callout(doc, "A Flexible Weekend in Samaná", passage, LIGHT_BLUE)
    add_question(doc, 27, "Where is the writer, and who is traveling with the writer?", answer_lines=1)
    add_question(doc, 28, "What are Elena and the writer doing right now?", answer_lines=1)
    add_question(doc, 29, "Write two arranged future plans from the blog.", answer_lines=1)
    add_question(doc, 30, "Why might the boat tour be canceled? Answer: _________________________________________________")
    add_question(doc, 31, "What is the backup plan, and which connector introduces the contrast? Answer: ______________________________")

    doc.add_page_break()

    add_section_heading(doc, "E", "WRITING: MY TRIP OR WEEKEND PLAN", 12)
    add_instruction(doc, "Write 70–90 words in two short paragraphs. Use your own ideas.")
    add_callout(
        doc,
        "Your writing must include",
        "• a clear introduction (place and purpose)   • one action happening now   • two arranged future plans with time expressions   • three useful adjectives   • one connector (because, however, therefore, or for example)",
        LIGHT_TEAL,
    )
    prompt = doc.add_paragraph()
    prompt.paragraph_format.space_before = Pt(8)
    prompt.paragraph_format.space_after = Pt(8)
    r = prompt.add_run("32. Write your blog post below. Give it a short title.")
    set_font(r, size=11, color=NAVY, bold=True)
    for _ in range(14):
        line = doc.add_paragraph()
        line.paragraph_format.space_before = Pt(0)
        line.paragraph_format.space_after = Pt(6)
        line.paragraph_format.line_spacing = 1.0
        lr = line.add_run("________________________________________________________________________________")
        set_font(lr, size=9, color="C8CDD3")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Word count: __________")
    set_font(r, size=10, color=MID_GRAY, bold=True)

    rubric = doc.add_table(rows=2, cols=6)
    set_table_geometry(rubric, [1560] * 6)
    set_table_borders(rubric, color=GRID, size=6, inside=True)
    labels = ["Intro", "Now", "2 plans", "3 adjectives", "Connector", "Accuracy"]
    for cell, label in zip(rubric.rows[0].cells, labels):
        set_cell_shading(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label)
        set_font(r, size=8.5, color=NAVY, bold=True)
    for cell in rubric.rows[1].cells:
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("___ / 2")
        set_font(r, size=9.5, color=INK)


def answer_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [1320, 8040])
    set_table_borders(table, color=GRID, size=6, inside=True)
    set_repeat_table_header(table.rows[0])
    set_row_cant_split(table.rows[0])
    for cell, text in zip(table.rows[0].cells, ("Item", "Answer / scoring guidance")):
        set_cell_shading(cell, NAVY)
        p = cell.paragraphs[0]
        r = p.add_run(text)
        set_font(r, size=9.5, color=WHITE, bold=True)
    for item, answer in rows:
        cells = table.add_row().cells
        set_row_cant_split(table.rows[-1])
        set_cell_width(cells[0], 1320)
        set_cell_width(cells[1], 8040)
        for cell in cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(item))
        set_font(r, size=9.2, color=NAVY, bold=True)
        p2 = cells[1].paragraphs[0]
        r2 = p2.add_run(answer)
        set_font(r2, size=9.2, color=INK)
    return table


def build_teacher_key(doc):
    doc.add_page_break()
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("TEACHER COPY")
    set_font(r, size=10, color=TEAL, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Answer Key & Scoring Guide")
    set_font(r, size=24, color=NAVY, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("Unit 1: Making Plans  |  50 points total")
    set_font(r, size=11.5, color=MID_GRAY)
    add_callout(doc, "Printing note", "Print pages 1–4 for the student. Keep the remaining pages as the teacher copy.", LIGHT_GRAY)

    add_section_heading(doc, "A", "WEATHER & MEANING", 8)
    answer_table(doc, [
        ("1–6", "1 cloudy; 2 windy; 3 warm; 4 foggy; 5 stormy; 6 dry — 1 point each."),
        ("7", "B. Watch a movie inside."),
        ("8", "B. It is cold, windy and cloudy."),
    ])

    add_section_heading(doc, "B", "GRAMMAR IN CONTEXT", 12)
    answer_table(doc, [
        ("9", "A. listens to — usual habit; third-person -s."),
        ("10", "B. are recording — action happening right now."),
        ("11", "A. visit — repeated routine with every Sunday."),
        ("12", "B. are you doing — present progressive question for an arranged Friday plan."),
        ("13", "B. are meeting — arranged future with a time and place."),
        ("14", "B. is going — arranged future this weekend."),
        ("15", "B. am not wearing — today's current situation."),
        ("16", "A. plays — usual Saturday routine; third-person -s."),
        ("17", "2 points: Harry isn't playing tennis at 2:00. He is studying for the science test. Award 1 point for the accurate negative correction and 1 for the true plan."),
        ("18", "2 points: any well-formed future-arrangement question, e.g., What time are you leaving tomorrow? Award 1 for question form and 1 for a future time expression."),
    ])

    add_section_heading(doc, "C", "INVITE & MAKE AN ARRANGEMENT", 8)
    answer_table(doc, [
        ("19", "Sorry, I can't."),
        ("20", "What a bummer!"),
        ("21", "Would you like to"),
        ("22", "That sounds great."),
        ("23", "How about"),
        ("24", "Let's meet"),
        ("25", "They are watching a movie on Sunday at 4:00."),
        ("26", "They are meeting outside the movie theater."),
    ])

    doc.add_page_break()
    add_section_heading(doc, "D", "READING: A WEEKEND BLOG", 10)
    answer_table(doc, [
        ("27", "The writer is in Samaná with the writer's cousin Elena."),
        ("28", "Elena is taking photos near the water, and the writer is writing the blog post at a café."),
        ("29", "Any two: visiting Los Haitises National Park; meeting cousins at 4:00 outside the museum; having lunch together on Sunday; returning home; visiting the museum if the boat tour is canceled."),
        ("30", "The tour might be canceled because the forecast says it may be stormy."),
        ("31", "The backup plan is to visit the museum. The connector is However."),
    ])

    add_section_heading(doc, "E", "WRITING RUBRIC", 12)
    add_instruction(doc, "Score each criterion from 0 to 2. Give brief, specific feedback.")
    table = doc.add_table(rows=1, cols=4)
    set_table_geometry(table, [2160, 2400, 2400, 2400])
    set_table_borders(table, color=GRID, size=6, inside=True)
    set_repeat_table_header(table.rows[0])
    set_row_cant_split(table.rows[0])
    headers = ["Criterion", "2 — Meets", "1 — Developing", "0 — Not shown"]
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, NAVY)
        p = cell.paragraphs[0]
        r = p.add_run(header)
        set_font(r, size=8.5, color=WHITE, bold=True)
    rubric_rows = [
        ("Introduction", "Clear place and purpose", "One detail is unclear", "No clear introduction"),
        ("Action now", "Accurate progressive action happening now", "Meaning is clear but form has an error", "Not included"),
        ("Future plans", "Two arranged plans with time expressions", "Only one complete plan or weak time detail", "No arranged future plan"),
        ("Adjectives", "Three appropriate adjectives", "One or two adjectives", "No useful adjectives"),
        ("Connector", "Connector is accurate and logical", "Connector is attempted but awkward", "No connector"),
        ("Accuracy & organization", "70–90 words; two readable paragraphs; errors do not impede meaning", "Some organization or accuracy problems", "Too incomplete to assess"),
    ]
    for row_data in rubric_rows:
        cells = table.add_row().cells
        set_row_cant_split(table.rows[-1])
        for idx, (cell, value) in enumerate(zip(cells, row_data)):
            set_cell_width(cell, [2160, 2400, 2400, 2400][idx])
            set_cell_margins(cell, top=90, bottom=90)
            p = cell.paragraphs[0]
            r = p.add_run(value)
            set_font(r, size=8.3, color=INK, bold=(idx == 0))

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run("Suggested interpretation: 45–50 = Strong mastery  |  38–44 = Ready with minor review  |  30–37 = Developing  |  Below 30 = Reteach priority")
    set_font(r, size=9.2, color=MID_GRAY, bold=True)


def main():
    doc = Document()
    configure_document(doc)
    build_student_pages(doc)
    build_teacher_key(doc)

    props = doc.core_properties
    props.title = "Unit 1: Making Plans — Evaluation Quiz"
    props.subject = "Weather, present tenses, future arrangements, invitations, reading and writing"
    props.author = "APEC English"
    props.keywords = "Unit 1, Making Plans, English evaluation, A2"

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
